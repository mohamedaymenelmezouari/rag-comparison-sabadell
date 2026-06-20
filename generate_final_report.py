"""
generate_final_report.py — Generates Capstone_Report_Final.docx using python-docx (Python only).
Run: python3 generate_final_report.py
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x00, 0x20, 0x54)
BLUE  = RGBColor(0x00, 0x6D, 0xFF)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY  = RGBColor(0x64, 0x74, 0x8B)
GREEN = RGBColor(0x05, 0x96, 0x69)
DARK  = RGBColor(0x1E, 0x29, 0x3B)

OUT_PATH = Path(__file__).parent / "Capstone_Report_Final.docx"


def set_cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def cell_para(cell, text: str, bold=False, color: RGBColor | None = None,
              align=WD_ALIGN_PARAGRAPH.LEFT, size: int = 9) -> None:
    para = cell.paragraphs[0]
    para.alignment = align
    run  = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
        run.bold = True


def body(doc: Document, text: str, indent: float = 0) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK


def bullet(doc: Document, text: str, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color


def bold_intro(doc: Document, label: str, rest: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.space_before = Pt(2)
    r1 = p.add_run(label)
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
    r2 = p.add_run(rest)
    r2.font.size = Pt(10.5); r2.font.color.rgb = DARK


def two_col_table(doc: Document, rows: list[tuple[str, str]],
                  h1: str = "", h2: str = "", widths=(7, 9.5)) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(widths[0])
    tbl.columns[1].width = Cm(widths[1])
    if h1:
        for i, hdr in enumerate([h1, h2]):
            set_cell_bg(tbl.rows[0].cells[i], "002054")
            cell_para(tbl.rows[0].cells[i], hdr, bold=True, color=WHITE, size=9)
    for ri, (c0, c1) in enumerate(rows, 1):
        bg = "F1F5F9" if ri % 2 == 0 else "FFFFFF"
        set_cell_bg(tbl.rows[ri].cells[0], bg)
        set_cell_bg(tbl.rows[ri].cells[1], bg)
        cell_para(tbl.rows[ri].cells[0], c0, bold=True, color=NAVY, size=9)
        cell_para(tbl.rows[ri].cells[1], c1, size=9)
    doc.add_paragraph()


def three_col_table(doc: Document, rows: list[tuple[str, str, str]],
                    h1: str = "", h2: str = "", h3: str = "",
                    widths=(5.5, 5.5, 5.5)) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(widths):
        tbl.columns[i].width = Cm(w)
    if h1:
        for i, hdr in enumerate([h1, h2, h3]):
            set_cell_bg(tbl.rows[0].cells[i], "002054")
            cell_para(tbl.rows[0].cells[i], hdr, bold=True, color=WHITE, size=9)
    for ri, (c0, c1, c2) in enumerate(rows, 1):
        bg = "F1F5F9" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate([c0, c1, c2]):
            set_cell_bg(tbl.rows[ri].cells[ci], bg)
            cell_para(tbl.rows[ri].cells[ci], val, size=9,
                      bold=(ci == 0), color=(NAVY if ci == 0 else None))
    doc.add_paragraph()


def four_col_table(doc: Document, rows: list[tuple], headers: list[str],
                   widths=(4, 4, 4, 4)) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(widths):
        tbl.columns[i].width = Cm(w)
    for i, hdr in enumerate(headers):
        set_cell_bg(tbl.rows[0].cells[i], "002054")
        cell_para(tbl.rows[0].cells[i], hdr, bold=True, color=WHITE, size=9)
    for ri, row_data in enumerate(rows, 1):
        bg = "F1F5F9" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            set_cell_bg(tbl.rows[ri].cells[ci], bg)
            cell_para(tbl.rows[ri].cells[ci], val, bold=(ci == 0),
                      color=(NAVY if ci == 0 else (GREEN if ci == 2 and ri > 1 else None)), size=9)
    doc.add_paragraph()


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build_report() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ══════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESADE CAPSTONE PROJECT 2026")
    r.font.size = Pt(11); r.font.color.rgb = GREY; r.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Traditional RAG vs. Graph RAG")
    r.font.size = Pt(26); r.font.color.rgb = NAVY; r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Policies and Product Catalogs in a Regulated Banking Environment")
    r.font.size = Pt(13); r.font.color.rgb = BLUE

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Banc Sabadell  |  Final Assessment Report  |  June 2026")
    r.font.size = Pt(11); r.font.color.rgb = GREY

    doc.add_paragraph()
    doc.add_paragraph()

    tbl = doc.add_table(rows=2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    names = ["Mohamed Aymen Elmezouari", "Abhay Mathahalli", "Warren Liu"]
    roles = [
        "Team Lead / AI Architecture\nPipeline Implementation",
        "Knowledge Graph Design\nNeo4j Schema and Graph RAG",
        "Corpus Curation / Evaluation\nBenchmarking and Reporting",
    ]
    for i in range(3):
        set_cell_bg(tbl.rows[0].cells[i], "002054")
        cell_para(tbl.rows[0].cells[i], names[i], bold=True, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        set_cell_bg(tbl.rows[1].cells[i], "F1F5F9")
        cell_para(tbl.rows[1].cells[i], roles[i], color=GREY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=8)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESADE Business School  |  Banc Sabadell Capstone Project 2026")
    r.font.size = Pt(9); r.font.color.rgb = GREY; r.italic = True

    page_break(doc)

    # ══════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════
    heading(doc, "1. Executive Summary")

    body(doc,
         "This report presents the final findings of the ESADE Capstone Project 2026 commissioned "
         "by Banc Sabadell. The project compares two AI retrieval architectures empirically: "
         "Traditional vector-based RAG and Graph RAG, which uses knowledge graphs to represent "
         "relationships between regulatory entities. The evaluation was designed around the specific "
         "requirements of a regulated financial institution, including GDPR compliance, "
         "explainability under MiFID II, and hallucination control for compliance-critical workflows.")

    body(doc,
         "Our central research question asked under what query complexity conditions Graph RAG "
         "outperforms Traditional RAG, and at what operational cost. After conducting a "
         "benchmarking study across 75 annotated query pairs stratified into four complexity "
         "tiers, we are able to answer this question with high empirical confidence.")

    heading(doc, "Key Findings", level=2)

    bold_intro(doc, "Accuracy on complex queries. ",
               "Graph RAG achieves significantly higher accuracy on Tier 3 multi-hop relational "
               "queries, outperforming Traditional RAG by 21 percentage points (89% compared to "
               "68%). This confirms the theoretical advantage of structured graph traversal for "
               "regulatory reasoning tasks that require connecting multiple entities across "
               "different directives.")

    bold_intro(doc, "Hallucination resistance. ",
               "Graph RAG reduces hallucination risk by 57 percent, from 21 percent with "
               "Traditional RAG down to 9 percent. In a compliance environment where a "
               "fabricated regulatory citation can constitute a material risk, this is the "
               "most practically significant finding of the study.")

    bold_intro(doc, "Speed advantage for simple queries. ",
               "Traditional RAG maintains a latency advantage of 0.8 seconds on average "
               "(2.3s compared to 3.1s) and reaches near-parity accuracy on Tier 1 simple "
               "factual queries, making it the more cost-effective choice for high-volume "
               "lookup workloads where complex reasoning is not required.")

    bold_intro(doc, "Statistical conclusion. ",
               "The null hypothesis of no significant performance difference between "
               "architectures is rejected for Tier 2, Tier 3, and Tier 4 queries. "
               "It cannot be rejected for Tier 1.")

    heading(doc, "Recommendation", level=2)

    body(doc,
         "We recommend a hybrid query-routing architecture for Banc Sabadell. A lightweight "
         "query complexity classifier, requiring less than 100 milliseconds of inference time, "
         "would dispatch Tier 1 queries to Traditional RAG for speed while directing Tier 2 "
         "through Tier 4 queries to Graph RAG for accuracy and explainability. All "
         "compliance-critical workflows, regardless of query tier, should route exclusively "
         "through Graph RAG because of its auditable traversal paths and substantially lower "
         "hallucination rate. This hybrid approach is estimated to capture 94 percent of "
         "Graph RAG's accuracy benefit while keeping average response latency around 2.5 "
         "seconds across the full query mix.")

    body(doc,
         "The team resolved its primary data constraint by constructing a corpus from publicly "
         "mandatory EU banking disclosures rather than internal Sabadell documents. This pivot "
         "produced a benchmark that is legally accessible, reproducible by future teams, and "
         "directly representative of the regulatory landscape Sabadell navigates daily.")

    page_break(doc)

    # ══════════════════════════════════════════════════════
    # 2. TEAM AND PROJECT CONTEXT
    # ══════════════════════════════════════════════════════
    heading(doc, "2. Team and Project Context")

    heading(doc, "2.1 The Team", level=2)
    three_col_table(doc, [
        ("Mohamed Aymen Elmezouari",
         "Team lead / AI and ML architecture",
         "Pipeline implementation, Groq API integration, Traditional RAG pipeline, Streamlit prototype"),
        ("Abhay Mathahalli",
         "Knowledge graph design / Neo4j schema",
         "Graph RAG pipeline construction, entity extraction using spaCy and LLM-assisted NER"),
        ("Warren Liu",
         "Corpus curation / evaluation methodology",
         "Metrics design, Ragas integration, 75-pair annotated test set, final reporting"),
    ], h1="Team Member", h2="Role", h3="Responsibilities", widths=(4.5, 4.5, 7.5))

    heading(doc, "2.2 Industry Partner", level=2)
    body(doc,
         "Banc Sabadell is one of Spain's leading financial institutions, headquartered in "
         "Sant Cugat del Valles, with over 230 billion euros in assets under management. "
         "The institution commissioned this capstone project to evaluate emerging AI retrieval "
         "architectures for potential deployment in compliance, customer service, and internal "
         "knowledge management workflows. Banc Sabadell's IT and Operations teams operate a "
         "predominantly Python-based technology stack, which shaped the project's technical "
         "decisions throughout.")

    heading(doc, "2.3 Project Origin and Data Constraint", level=2)
    body(doc,
         "The project was originally scoped to use internal Sabadell HR FAQs and product "
         "catalogs. Following discussions with the Sabadell mentor, the team learned that no "
         "internal proprietary data could be shared. Rather than treating this as a limitation, "
         "the team constructed an equivalent corpus from publicly available EU regulatory "
         "documents that Sabadell is legally required to comply with. This approach produced a "
         "benchmark that is more rigorous than a synthetic dataset would have been, because "
         "every document in the corpus reflects real compliance obligations the institution "
         "faces. The corpus strategy has been consistently praised as a strength of the project.")

    # ══════════════════════════════════════════════════════
    # 3. PROBLEM STATEMENT
    # ══════════════════════════════════════════════════════
    heading(doc, "3. Problem Statement and Research Question")

    heading(doc, "3.1 The Business Problem", level=2)
    body(doc,
         "Banks such as Sabadell manage thousands of pages of regulatory documents, internal "
         "policies, product manuals, and compliance guidelines. Employees and automated systems "
         "query this corpus constantly for compliance checks, customer queries, product "
         "recommendations, and audit trails. The quality of these retrieval systems has direct "
         "consequences for regulatory risk, customer satisfaction, and operational efficiency.")

    body(doc,
         "Current approaches that rely on keyword search or basic vector retrieval struggle "
         "in several important situations. When a query requires reasoning across multiple "
         "documents simultaneously, such systems tend to surface relevant fragments without "
         "the ability to connect them systematically. When an answer depends on the relationship "
         "between entities rather than the content of a single passage, vector similarity "
         "returns semantically close text without capturing that structural dependency. "
         "In a regulated environment where source citations must be auditable and where a "
         "fabricated answer carries compliance risk, these limitations become operationally "
         "significant.")

    heading(doc, "3.2 Why Graph RAG is Theoretically Appealing for Banking Compliance", level=2)
    body(doc,
         "Traditional RAG retrieves the most semantically similar text chunks to a query and "
         "passes them as context to a language model. This works well for factual lookups but "
         "cannot natively reason across entity relationships, which is a structural constraint "
         "that becomes visible precisely on the multi-hop regulatory queries that compliance "
         "officers encounter most often.")

    body(doc,
         "Graph RAG extends the retrieval step by encoding a knowledge graph, which is a "
         "network of typed entities and relationships that captures the structural connections "
         "between regulatory concepts. In a banking compliance context this means the system "
         "can explicitly represent that Article 17 of MiFID II governs best execution "
         "obligations, and that those obligations apply to retail investment products. "
         "Rather than retrieving the most semantically similar text chunk, Graph RAG traverses "
         "this graph to assemble contextually connected evidence. The traversal path itself "
         "serves as an auditable reasoning trail, which has significant explainability "
         "advantages under MiFID II and EBA supervisory expectations.")

    heading(doc, "3.3 Research Questions", level=2)
    body(doc,
         "The primary research question guiding this project was: under what query complexity "
         "conditions does Graph RAG outperform Traditional RAG for financial regulatory document "
         "retrieval, and what are the trade-offs in accuracy, hallucination rate, latency, "
         "and maintainability?")

    body(doc, "Four secondary questions shaped the experimental design:")
    for item in [
        "How does each architecture perform on simple factual queries relative to multi-hop "
        "reasoning queries, and does the performance gap follow a predictable pattern?",
        "What is the build and maintenance overhead of a knowledge graph compared to a vector "
        "index, and how does that overhead scale with corpus size?",
        "What GDPR and data protection implications arise from each architecture's approach to "
        "data handling, entity linking, and graph traversal?",
        "Under what organisational conditions should a financial institution prefer one "
        "architecture over the other, and can both be combined in a single deployment?",
    ]:
        bullet(doc, item)

    # ══════════════════════════════════════════════════════
    # 4. TECHNICAL ARCHITECTURE
    # ══════════════════════════════════════════════════════
    heading(doc, "4. Technical Architecture")

    heading(doc, "4.1 Traditional RAG Pipeline", level=2)
    body(doc,
         "The Traditional RAG pipeline uses dense vector retrieval. Documents are parsed with "
         "PyMuPDF and split into 500-token chunks with a 50-token overlap using LangChain's "
         "RecursiveCharacterTextSplitter. Each chunk is embedded with the "
         "sentence-transformers/all-MiniLM-L6-v2 model, which runs entirely locally without "
         "requiring an API call, and the resulting vectors are persisted in a ChromaDB store. "
         "At query time, the top-k chunks by cosine similarity are retrieved and passed as "
         "context to Llama 3.3-70B through the Groq API, achieving sub-2.5-second average "
         "latency.")

    body(doc,
         "ChromaDB was selected over cloud-hosted alternatives because it runs fully locally, "
         "ensuring that no regulatory text leaves the local environment. This is a hard "
         "requirement given GDPR's data minimisation obligations. The all-MiniLM-L6-v2 "
         "embedding model was chosen for its strong performance on domain-general semantic "
         "tasks, its open-source licence, and its zero API cost. The 500-token chunk size "
         "with 50-token overlap was calibrated through preliminary testing to preserve "
         "sufficient regulatory context, typically covering a full article or sub-section, "
         "while keeping retrieval precision high.")

    heading(doc, "4.2 Graph RAG Pipeline", level=2)
    body(doc,
         "The Graph RAG pipeline introduces a structured knowledge layer between document "
         "ingestion and answer generation. During indexing, spaCy with a custom financial "
         "NER model recognises eight entity types across the corpus: Regulation, Article, "
         "Product, Obligation, Role, Risk Category, Threshold, and Date. LLM-assisted "
         "extraction then identifies six relationship types connecting these entities: "
         "GOVERNS, REQUIRES, APPLIES TO, REFERENCES, SUPERSEDES, and EXEMPTS. These entity "
         "types and relationship types were defined through a deep reading of the corpus "
         "rather than imported from a generic ontology, so they reflect the actual conceptual "
         "structure of EU banking regulation.")

    body(doc,
         "The resulting graph is stored in Neo4j Community Edition 5.x, which provides a "
         "native Cypher query interface well suited to traversing regulatory entity networks. "
         "At query time, key terms are extracted from the incoming query and matched against "
         "the graph through Cypher subgraph queries. A relevant subgraph is assembled, "
         "enriched with co-located text chunks for factual grounding, and passed as context "
         "to Llama 3.3-70B for answer generation. The full traversal path is logged at the "
         "node and edge level, providing a provenance record that satisfies audit and "
         "explainability requirements.")

    body(doc,
         "The most important improvement from the mid-term prototype to the final evaluation "
         "was the increase in relationship density. The mid-term graph contained 22 "
         "relationships across 5,235 entities, which amounted to a coverage rate of 0.4 "
         "percent and meant that graph traversal could not reliably surface multi-hop "
         "reasoning paths. The entity matching threshold was lowered from three characters "
         "to two, with an acronym-specific whitelist added for terms such as AML, CDD, PEP, "
         "and KYC. Relationship extraction was refactored to use LLM-assisted extraction "
         "across sentence pairs, which increased relationship density from 22 to 847 "
         "relationships, a 38.5-fold improvement that materially changed the performance "
         "comparison in the final evaluation.")

    heading(doc, "4.3 Technology Stack", level=2)
    three_col_table(doc, [
        ("Language",            "Python 3.11",                          "All pipelines, evaluation, and prototype UI (Python-only stack)"),
        ("LLM",                 "Llama 3.3-70B via Groq API",           "Answer generation for both pipelines, sub-2-second latency"),
        ("Embeddings",          "all-MiniLM-L6-v2 (HuggingFace local)","Document and query embedding with no API cost and full data locality"),
        ("Vector Store",        "ChromaDB 0.4.x (local instance)",      "Vector similarity search with top-k cosine retrieval"),
        ("Graph Database",      "Neo4j Community 5.x with APOC",        "Knowledge graph storing 5,235 entities and 847 relationships"),
        ("NLP and NER",         "spaCy en_core_web_sm with custom rules","Eight entity types and six relationship types for the financial domain"),
        ("Evaluation",          "Ragas 0.1.x with GPT-4o judge",        "Automated metrics using an independent judge to avoid circularity bias"),
        ("Frontend",            "Streamlit with Plotly (Python)",        "Interactive comparison prototype built on a Python-only stack"),
        ("Orchestration",       "LangChain 0.2.x",                      "Document processing, text splitting, and prompt management"),
    ], h1="Component", h2="Technology", h3="Notes", widths=(4, 5, 7.5))

    # ══════════════════════════════════════════════════════
    # 5. EXPERIMENTAL DESIGN
    # ══════════════════════════════════════════════════════
    heading(doc, "5. Experimental Design and Corpus")

    heading(doc, "5.1 Document Corpus", level=2)
    body(doc,
         "The evaluation corpus comprises nine documents totalling approximately 750 pages. "
         "All documents are publicly available mandatory disclosures that directly represent "
         "the compliance landscape Banc Sabadell navigates daily. Three sample text files "
         "were included to ensure broad coverage of key regulatory concepts in a form that "
         "can be loaded without PDF access restrictions.")
    three_col_table(doc, [
        ("EBA AML/CFT Guidelines 2021",         "EBA/GL/2021/02",  "87 pages covering AML, CDD, EDD, and SAR procedures"),
        ("BIS Basel III Framework",              "BIS 2017",        "63 pages covering CET1, leverage ratios, LCR, and NSFR"),
        ("MiFID II Directive 2014/65/EU",        "EUR-Lex",         "90 pages covering investor protection and best execution"),
        ("GDPR Regulation 2016/679/EU",          "EUR-Lex",         "54 pages covering data protection and breach notification"),
        ("Banc Sabadell Pillar III Report 2025", "BSAB-PIR-2025",  "210 pages covering risk disclosures and capital adequacy"),
        ("Banc Sabadell Annual Report 2025",     "BSAB-AR-2025",   "Approximately 350 pages covering financial performance and products"),
        ("EBA AML Guidelines sample text",       "Supplementary",   "15 pages of key AML concepts for corpus breadth"),
        ("Basel III sample text",                "Supplementary",   "12 pages of core prudential concepts"),
        ("GDPR Banking sample text",             "Supplementary",   "14 pages of GDPR banking applications"),
    ], h1="Document", h2="Source", h3="Description", widths=(5.5, 3, 8))

    heading(doc, "5.2 Query Complexity Tiers", level=2)
    body(doc,
         "The 75 annotated query pairs are stratified across four complexity tiers. Ground-truth "
         "answers were produced by the research team, with 20 percent of the set "
         "cross-validated by GPT-4o as a second annotator. The tier structure was designed so "
         "that each successive tier requires a qualitatively more complex form of reasoning, "
         "allowing the analysis to identify precisely where the performance gap between "
         "architectures begins to emerge.")
    three_col_table(doc, [
        ("Tier 1: Simple Factual",      "20 pairs",
         "Single-document, single-section lookup. Example: What is the minimum CET1 capital ratio under Basel III?"),
        ("Tier 2: Single-hop Relational","20 pairs",
         "Combines information from two linked entities. Example: What CDD obligations apply specifically to Politically Exposed Persons?"),
        ("Tier 3: Multi-hop Relational", "20 pairs",
         "Requires traversing three or more entity relationships. Example: Which products require both MiFID II suitability assessment and AML enhanced due diligence?"),
        ("Tier 4: Adversarial",          "15 pairs",
         "Designed to elicit hallucination. Example: What does Article 47 of GDPR say about banking? This article does not address banking directly."),
    ], h1="Tier", h2="Size", h3="Description and Example", widths=(4, 2.5, 10))

    heading(doc, "5.3 Evaluation Metrics", level=2)
    body(doc,
         "Scoring was performed by GPT-4o as an independent judge model, which eliminates "
         "the circularity bias identified in the mid-term evaluation where the same "
         "Groq/Llama model was used for both generation and scoring. Six metrics were "
         "computed for each query pair.")
    three_col_table(doc, [
        ("Accuracy",               "LLM-as-judge rated on a 1 to 100 scale",     "Both pipelines"),
        ("Faithfulness",           "Proportion of claims grounded in retrieved context", "Both pipelines"),
        ("Hallucination Risk",     "Proportion of claims not traceable to source documents", "Both pipelines"),
        ("Source Citation Quality","Precision and recall of document citations",  "Both pipelines"),
        ("Answer Quality",         "Overall quality covering completeness, clarity, and relevance", "Both pipelines"),
        ("Latency",                "Wall-clock time per query in seconds, logged in Python", "Both pipelines"),
    ], h1="Metric", h2="Measurement Method", h3="Scope", widths=(4.5, 6.5, 5.5))

    # ══════════════════════════════════════════════════════
    # 6. RESULTS AND EVALUATION
    # ══════════════════════════════════════════════════════
    heading(doc, "6. Results and Evaluation")

    heading(doc, "6.1 Overall Summary", level=2)
    body(doc,
         "Graph RAG won on 52 of the 75 query pairs (69.3 percent), Traditional RAG won on "
         "18 (24.0 percent), and five pairs resulted in a tie. The distribution of wins is "
         "not uniform across tiers: Traditional RAG won predominantly on Tier 1 queries, "
         "while Graph RAG accumulated its wins almost entirely on Tier 2 through Tier 4.")
    three_col_table(doc, [
        ("Graph RAG wins",       "52 out of 75 (69.3%)", "Advantage concentrated in Tiers 2 to 4"),
        ("Traditional RAG wins", "18 out of 75 (24.0%)", "Concentrated in Tier 1 simple factual queries"),
        ("Ties",                 "5 out of 75 (6.7%)",   "Simple factual queries with single-source answers"),
    ], h1="Outcome", h2="Count", h3="Notes", widths=(5, 4, 7.5))

    heading(doc, "6.2 Metric Comparison Across All Tiers", level=2)
    three_col_table(doc, [
        ("Accuracy",               "79%",  "87%  (+8 pp)"),
        ("Faithfulness",           "82%",  "91%  (+9 pp)"),
        ("Source Citation Quality","74%",  "88%  (+14 pp)"),
        ("Answer Quality",         "81%",  "89%  (+8 pp)"),
        ("Hallucination Risk",     "21%",  "9%   (reduction of 57%)"),
        ("Average Latency",        "2.3s", "3.1s (overhead of 0.8 seconds)"),
    ], h1="Metric", h2="Traditional RAG", h3="Graph RAG", widths=(5.5, 5.5, 5.5))

    heading(doc, "6.3 Results by Query Complexity Tier", level=2)
    body(doc,
         "The per-tier breakdown reveals the central insight of this study. The performance "
         "gap between architectures grows consistently with query complexity, which confirms "
         "the theoretical prediction that graph-based retrieval confers its greatest advantage "
         "precisely where vector similarity search is weakest.")
    four_col_table(doc, [
        ("Tier 1: Simple Factual (20 pairs)",       "91%", "89% (Traditional leads by 2 pp)", "12% vs 6%"),
        ("Tier 2: Single-hop Relational (20 pairs)","84%", "90% (Graph leads by 6 pp)",       "18% vs 8%"),
        ("Tier 3: Multi-hop Relational (20 pairs)", "68%", "89% (Graph leads by 21 pp)",      "29% vs 9%"),
        ("Tier 4: Adversarial (15 pairs)",          "71%", "80% (Graph leads by 9 pp)",       "31% vs 13%"),
    ], headers=["Tier", "Traditional RAG Accuracy", "Graph RAG Accuracy", "Hallucination (Trad vs Graph)"],
    widths=[5, 3.5, 4, 4])

    heading(doc, "6.4 Analytical Interpretation", level=2)
    body(doc,
         "On Tier 1 queries, Traditional RAG holds a marginal advantage of two percentage "
         "points. This difference falls within the noise of the scoring methodology and does "
         "not represent a meaningful practical distinction. Both architectures perform "
         "equivalently on straightforward factual lookups from a single document.")

    body(doc,
         "On Tier 3 multi-hop relational queries, the 21 percentage point advantage for Graph "
         "RAG is statistically meaningful and operationally significant. These are precisely "
         "the queries that compliance officers and risk managers encounter most frequently, "
         "where answering correctly requires connecting obligations across multiple regulatory "
         "frameworks simultaneously. Vector similarity search surfaces relevant fragments but "
         "cannot systematically assemble the relational chain, whereas graph traversal follows "
         "the typed relationships between entities to construct a contextually complete answer.")

    body(doc,
         "The 57 percent reduction in hallucination risk from 21 percent to 9 percent is the "
         "finding that is most directly relevant to Banc Sabadell's deployment decision. A "
         "compliance system that fabricates regulatory citations, even occasionally, creates "
         "material risk. The grounded subgraph context in Graph RAG constrains the language "
         "model's generation to information that is explicitly supported by the retrieved "
         "evidence, which is why the hallucination rate drops so substantially.")

    # ══════════════════════════════════════════════════════
    # 7. DECISION FRAMEWORK
    # ══════════════════════════════════════════════════════
    heading(doc, "7. Deployment Decision Framework for Banc Sabadell")

    heading(doc, "7.1 Primary Recommendation: Hybrid Query-Routing Architecture", level=2)
    body(doc,
         "Neither architecture dominates across all query types, which means the optimal "
         "deployment for Banc Sabadell is not a choice between the two but a combination of "
         "both. We recommend a hybrid query-routing system in which a lightweight BERT-based "
         "query complexity classifier, fine-tuned on the 75-pair annotated set and requiring "
         "less than 100 milliseconds of inference time, routes each incoming query to the "
         "most appropriate pipeline. This adds negligible overhead while capturing the "
         "strengths of both architectures across the full query mix.")

    body(doc,
         "The routing logic is as follows. Tier 1 queries go to Traditional RAG because it "
         "is faster and equally accurate for simple factual lookups. Tier 2 through Tier 4 "
         "queries go to Graph RAG because its accuracy advantage grows with complexity. Any "
         "query that is flagged as compliance-critical, regardless of its complexity tier, "
         "routes to Graph RAG because the auditable traversal path and lower hallucination "
         "rate are non-negotiable in that context.")

    heading(doc, "7.2 When to Choose Traditional RAG", level=2)
    body(doc,
         "Traditional RAG is the right choice when the query mix consists predominantly of "
         "Tier 1 simple factual lookups, when a hard sub-2-second SLA exists such as in a "
         "real-time customer-facing chatbot, when the knowledge base updates frequently and "
         "the graph schema maintenance overhead would outweigh the accuracy gains, or when "
         "infrastructure constraints preclude running a Neo4j instance. Suitable use cases "
         "include customer FAQ chatbots, product information lookup, and single-document "
         "question-answering.")

    heading(doc, "7.3 When to Choose Graph RAG", level=2)
    body(doc,
         "Graph RAG is the right choice when queries require multi-hop regulatory reasoning "
         "across multiple directives simultaneously, when explainability and an audit trail "
         "are legally required under MiFID II or EBA supervisory guidance, when hallucination "
         "risk must be minimised because the output will inform a compliance decision, or "
         "when cross-framework queries span Basel III, MiFID II, GDPR, and AML obligations "
         "at the same time. Suitable use cases include compliance officer decision support, "
         "regulatory audit preparation, and impact analysis for new directives such as DORA "
         "and the EU AI Act.")

    heading(doc, "7.4 Phased Deployment Roadmap", level=2)
    three_col_table(doc, [
        ("Phase 1: Q3 2026",
         "Traditional RAG in production",
         "Deploy vector RAG for internal FAQ and product catalog queries. Collect live query "
         "logs to train the complexity classifier."),
        ("Phase 2: Q4 2026",
         "Graph RAG for compliance workflows",
         "Deploy Graph RAG for compliance officer query support and regulatory audit use cases, "
         "targeting MiFID II suitability queries and AML EDD decisions."),
        ("Phase 3: Q1 2027",
         "Hybrid routing architecture",
         "Implement the query complexity classifier and a unified API layer that routes Tier 1 "
         "to Traditional RAG and Tier 2 through Tier 4 to Graph RAG."),
        ("Phase 4: Q2 2027",
         "Graph expansion and Spanish corpus",
         "Extend the knowledge graph with Banco de Espana circulars and CNMV guidance. "
         "Begin evaluating the hybrid architecture on Spanish-language documents."),
    ], h1="Phase", h2="Initiative", h3="Description", widths=(3, 4.5, 9))

    # ══════════════════════════════════════════════════════
    # 8. GDPR AND COMPLIANCE ANALYSIS
    # ══════════════════════════════════════════════════════
    heading(doc, "8. GDPR and Regulatory Compliance Analysis")

    body(doc,
         "Both architectures were assessed against GDPR obligations and EBA AI governance "
         "expectations. The table below summarises the compliance implications identified "
         "for each architecture and the mitigations the team recommends before production "
         "deployment.")

    four_col_table(doc, [
        ("Data Minimisation",
         "Compliant: retrieves only the top-k chunks relevant to the query",
         "Requires scope-limited subgraph traversal to avoid surfacing unrelated entities",
         "Both architectures can comply with proper query scoping"),
        ("Explainability under MiFID II",
         "Limited: no reasoning path is exposed to the end user",
         "Strong: the full traversal path is logged at node and edge level",
         "Graph RAG is preferred for any workflow that must justify its output to regulators"),
        ("Right to Erasure",
         "Simple: deleting a chunk from the vector store removes the data",
         "Complex: deletion requires cascading removal of nodes and all relational dependencies",
         "A cascading deletion policy must be implemented in the graph before production"),
        ("Audit Trail",
         "Weak: there is no explicit log of how an answer was constructed",
         "Strong: node and edge level provenance is recorded for every query",
         "Graph RAG is required for any workflow subject to audit under EBA or ECB supervision"),
        ("Access Control",
         "Collection-level access control is sufficient",
         "Edge and path-level RBAC is needed to prevent indirect exposure of restricted data",
         "Path-level RBAC must be implemented before production deployment of Graph RAG"),
        ("Hallucination Risk",
         "Medium at 21 percent across the test set",
         "Low at 9 percent, a reduction of 57 percent relative to Traditional RAG",
         "Graph RAG is preferred for compliance-critical workflows"),
        ("Profiling and Inference Risk",
         "Low: no cross-entity inference takes place during retrieval",
         "Medium: graph traversal can enable inference of sensitive attributes from relationships",
         "Inference on sensitive attributes must be restricted and derived features must be audited"),
    ], headers=["Dimension", "Traditional RAG", "Graph RAG", "Recommendation"],
    widths=[3.5, 3.8, 3.8, 5.4])

    # ══════════════════════════════════════════════════════
    # 9. CONCLUSIONS AND FUTURE WORK
    # ══════════════════════════════════════════════════════
    heading(doc, "9. Conclusions and Future Work")

    heading(doc, "9.1 Conclusions", level=2)
    body(doc,
         "This project produced a rigorous, empirically grounded comparison of Traditional "
         "RAG and Graph RAG on a banking regulatory corpus of nine documents and approximately "
         "750 pages, evaluated on 75 annotated query pairs across four complexity tiers. The "
         "primary research question is answered conclusively.")

    body(doc,
         "Graph RAG outperforms Traditional RAG on multi-hop relational queries by 21 "
         "percentage points, on adversarial queries by 9 percentage points, and on "
         "hallucination resistance by 57 percent. Traditional RAG maintains a speed "
         "advantage of 0.8 seconds and near-parity accuracy on simple factual queries. "
         "The recommended hybrid architecture captures 94 percent of Graph RAG's accuracy "
         "improvement while maintaining an average latency of approximately 2.5 seconds, "
         "which is operationally viable for Banc Sabadell's compliance and customer service "
         "workflows.")

    body(doc,
         "The constraint that prevented access to internal Sabadell data, which appeared "
         "to be the project's greatest risk at inception, became one of its defining "
         "strengths. The resulting benchmark is reproducible, legally accessible, and "
         "directly representative of real compliance obligations. Future Sabadell teams can "
         "extend it without any data access concerns and use it as the foundation for "
         "ongoing architecture evaluation.")

    heading(doc, "9.2 Future Work", level=2)
    body(doc, "Six directions for future development are identified in order of priority.")
    for item in [
        "Implement the hybrid query-routing architecture with a BERT-based complexity "
        "classifier trained on the annotated set, and measure its latency and accuracy "
        "trade-off against single-architecture deployment in a production environment.",
        "Extend the knowledge graph with Spanish-language regulatory documents, including "
        "Banco de Espana circulars and CNMV guidance, to support Sabadell's domestic "
        "compliance workflows.",
        "Evaluate the implications of DORA and the EU AI Act for both retrieval "
        "architectures, particularly regarding model documentation, risk classification, "
        "and human oversight requirements.",
        "Conduct a human evaluation study with Sabadell compliance officers using a real "
        "internal query set, which would validate the benchmark results against practitioner "
        "judgement and reveal query patterns not captured in the annotated test set.",
        "Explore a graph update pipeline that handles regulatory document amendments and "
        "new directive publications without requiring full re-indexing of the knowledge graph.",
        "Investigate retrieval-augmented fine-tuning (RAFT) as a third architecture for "
        "comparison, which trains the language model on the retrieval corpus directly and "
        "may offer accuracy advantages for highly specialised regulatory domains.",
    ]:
        bullet(doc, item)

    # ══════════════════════════════════════════════════════
    # 10. REFERENCES
    # ══════════════════════════════════════════════════════
    heading(doc, "10. References")

    heading(doc, "Academic and Technical Sources", level=2)
    for i, ref in enumerate([
        "Lewis, P., Perez, E., Piktus, A., and others. (2020). Retrieval-Augmented "
        "Generation for Knowledge-Intensive NLP Tasks. Advances in Neural Information "
        "Processing Systems, 33, 9459-9474.",
        "Edge, D., Trinh, H., Cheng, N., and others. (2024). From Local to Global: "
        "A Graph RAG Approach to Query-Focused Summarization. Microsoft Research. "
        "arXiv: 2404.16130.",
        "Pan, S., Luo, L., Wang, Y., and others. (2024). Unifying Large Language Models "
        "and Knowledge Graphs: A Roadmap. IEEE Transactions on Knowledge and Data "
        "Engineering. arXiv: 2306.08302.",
        "Es, S., James, J., Espinosa-Anke, L., and Schockaert, S. (2023). RAGAS: "
        "Automated Evaluation of Retrieval Augmented Generation. arXiv: 2309.15217.",
        "Reimers, N. and Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using "
        "Siamese BERT-Networks. Proceedings of EMNLP 2019.",
        "Huang, Y. and Huang, J. (2026). A Survey on Retrieval-Augmented Text Generation "
        "for Large Language Models. ACM Computing Surveys 2026.",
    ], 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(ref)
        r.font.size = Pt(9.5); r.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(3)

    heading(doc, "Regulatory Sources", level=2)
    for i, ref in enumerate([
        "European Banking Authority (2021). EBA Guidelines on Anti-Money Laundering and "
        "Counter-Terrorist Financing (EBA/GL/2021/02). Available at eba.europa.eu.",
        "Bank for International Settlements (2017). Basel III: Finalising Post-Crisis "
        "Reforms. Available at bis.org.",
        "European Parliament and Council (2014). Directive 2014/65/EU on Markets in "
        "Financial Instruments (MiFID II). Available at EUR-Lex.",
        "European Parliament and Council (2016). Regulation (EU) 2016/679: General Data "
        "Protection Regulation. Available at EUR-Lex.",
        "Banc Sabadell (2025). Pillar III Risk Disclosures 2025 and Consolidated Annual "
        "Financial Report 2025. Available at grupbancsabadell.com.",
    ], 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(ref)
        r.font.size = Pt(9.5); r.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(3)

    page_break(doc)

    # ══════════════════════════════════════════════════════
    # APPENDIX A
    # ══════════════════════════════════════════════════════
    heading(doc, "Appendix A: Detailed Technology Stack")

    three_col_table(doc, [
        ("Language",              "Python 3.11",
         "All pipelines, evaluation framework, and Streamlit prototype are Python-only"),
        ("LLM",                   "Llama 3.3-70B via Groq API",
         "Primary model; Mistral 7B via Ollama serves as a fallback for offline use"),
        ("Embedding model",       "all-MiniLM-L6-v2 (HuggingFace local)",
         "Runs locally with no API cost, satisfying GDPR data minimisation obligations"),
        ("Vector store",          "ChromaDB 0.4.x (persistent local instance)",
         "Full data locality; no regulatory text leaves the local environment"),
        ("Graph database",        "Neo4j Community 5.x with APOC",
         "5,235 entities and 847 relationships with a native Cypher query interface"),
        ("Pipeline orchestration","LangChain 0.2.x",
         "Document processing, text splitting, and prompt construction"),
        ("Evaluation",            "Ragas 0.1.x with GPT-4o as independent judge",
         "Independent scoring model eliminates the circularity bias from the mid-term prototype"),
        ("Document parsing",      "PyMuPDF (fitz)",
         "PDF text extraction with page-level metadata preserved"),
        ("NLP and NER",           "spaCy en_core_web_sm with custom financial rules",
         "Eight entity types and six relationship types defined from corpus reading"),
        ("Frontend prototype",    "Streamlit with Plotly and Pandas",
         "Interactive comparison prototype built entirely in Python"),
    ], h1="Component", h2="Technology", h3="Notes", widths=(4, 5, 7.5))

    heading(doc, "Appendix B: Mid-Term to Final Improvements")

    three_col_table(doc, [
        ("Entity matching threshold",
         "Three-character minimum caused misses on AML, CDD, and PEP",
         "Lowered to two characters with an acronym whitelist added"),
        ("Relationship density",
         "22 relationships across 5,235 entities at mid-term (0.4 percent coverage)",
         "847 relationships (16.2 percent) through LLM-assisted sentence-pair extraction"),
        ("Evaluation judge model",
         "Groq/Llama used for both generation and scoring, introducing circularity bias",
         "GPT-4o used as independent judge model for the final evaluation"),
        ("Test set size",
         "20 annotated pairs at mid-term",
         "75 pairs with 20 percent cross-validated by a second annotator"),
        ("EBA AML document",
         "Only a sample text file was available due to a broken public URL",
         "Full 87-page EBA/GL/2021/02 PDF indexed in the final corpus"),
        ("Traditional RAG latency",
         "Six to eight seconds per query at mid-term due to inefficient indexing of large PDFs",
         "2.3 seconds on average after pre-filtering text-dense sections and refining chunking"),
        ("Prototype branding and logo",
         "SVG placeholder with approximate Banc Sabadell styling",
         "Official Banc Sabadell SVG logo (from public Wikipedia source) embedded directly"),
    ], h1="Area", h2="Mid-Term State", h3="Final State", widths=(4, 5.5, 7))

    p = doc.add_paragraph()
    r = p.add_run(
        "AI Disclosure: this report was drafted with AI assistance for document structuring, "
        "prose drafting, and diagram generation. All technical content, empirical results, "
        "design decisions, and analytical judgements reflect the work of the project team. "
        "AI tools were used as a productivity aid and not as a source of knowledge or evaluation."
    )
    r.font.size = Pt(9); r.font.color.rgb = GREY; r.italic = True
    p.paragraph_format.space_before = Pt(12)

    doc.save(str(OUT_PATH))
    print(f"Report saved to {OUT_PATH}")
    words = len(open(OUT_PATH, "rb").read()) // 6  # rough estimate
    print(f"File size: {OUT_PATH.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build_report()
